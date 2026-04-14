import streamlit as st
from datetime import datetime
import os
import re
import json
import time
import io
from openai import OpenAI
from dotenv import load_dotenv
from database import guardar_analisis, obtener_analisis

# ============================================
# LIBRERÍAS PARA EXTRACCIÓN DE PDF
# ============================================
try:
    from pypdf import PdfReader
    PDF_EXTRACTION = True
except:
    PDF_EXTRACTION = False
    st.warning("⚠️ pypdf no instalado. La extracción de PDF será limitada.")

# ============================================
# LIBRERÍAS PARA GENERAR WORD
# ============================================
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    WORD_AVAILABLE = True
except:
    WORD_AVAILABLE = False

# ============================================
# LIBRERÍAS PARA GENERAR PDF
# ============================================
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except:
    PDF_AVAILABLE = False

load_dotenv()

# ============================================
# CONFIGURACIÓN DE APIS
# ============================================

DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")

deepseek_client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com/v1") if DEEPSEEK_API_KEY else None
groq_client = OpenAI(api_key=GROQ_API_KEY, base_url="https://api.groq.com/openai/v1") if GROQ_API_KEY else None

# ============================================
# FUNCIONES DE EXTRACCIÓN DE TEXTO (REAL)
# ============================================

def extraer_texto_archivo(archivo):
    """Extrae texto REAL de PDF, DOCX, TXT, MD"""
    nombre = archivo.name
    contenido_bytes = archivo.getvalue()
    
    # Para PDF
    if nombre.lower().endswith('.pdf') and PDF_EXTRACTION:
        try:
            reader = PdfReader(io.BytesIO(contenido_bytes))
            texto = ""
            for page in reader.pages:
                texto += page.extract_text() or ""
            return texto
        except Exception as e:
            return f"[Error al leer PDF: {str(e)}]"
    
    # Para DOCX (extracción básica)
    elif nombre.lower().endswith('.docx') and WORD_AVAILABLE:
        try:
            doc = Document(io.BytesIO(contenido_bytes))
            texto = ""
            for para in doc.paragraphs:
                texto += para.text + "\n"
            return texto
        except:
            return contenido_bytes.decode('utf-8', errors='ignore')
    
    # Para TXT, MD y otros
    else:
        return contenido_bytes.decode('utf-8', errors='ignore')

def extraer_texto_multiple(archivos):
    """Extrae texto de múltiples archivos"""
    texto_completo = ""
    for a in archivos:
        texto_completo += f"\n\n--- {a.name} ---\n{extraer_texto_archivo(a)}"
    return texto_completo

# ============================================
# FUNCIONES DE GENERACIÓN DE WORD (REAL)
# ============================================

def generar_documento_word(contenido_markdown, titulo):
    """Genera un documento Word profesional con formato"""
    if not WORD_AVAILABLE:
        return None
    
    doc = Document()
    
    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
    
    # Título principal
    title = doc.add_heading(level=0)
    title_run = title.add_run(titulo.upper())
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Fecha
    fecha = doc.add_paragraph()
    fecha_run = fecha.add_run(f"Fecha de presentación: {datetime.now().strftime('%d de %B de %Y')}")
    fecha_run.font.size = Pt(11)
    fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    
    # Procesar contenido Markdown
    lineas = contenido_markdown.split('\n')
    for linea in lineas:
        if linea.startswith('# '):
            doc.add_heading(linea[2:], level=1)
        elif linea.startswith('## '):
            doc.add_heading(linea[3:], level=2)
        elif linea.startswith('### '):
            doc.add_heading(linea[4:], level=3)
        elif linea.strip().startswith('|') and '|' in linea:
            # Intentar crear tabla (simplificado)
            pass
        elif linea.strip():
            p = doc.add_paragraph(linea)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.size = Pt(11)
    
    # Guardar
    temp_file = f"propuesta_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(temp_file)
    return temp_file

# ============================================
# FUNCIONES DE GENERACIÓN DE PDF (REAL)
# ============================================

def generar_documento_pdf(contenido_markdown, titulo):
    """Genera un documento PDF profesional con formato"""
    if not PDF_AVAILABLE:
        return None
    
    temp_file = f"propuesta_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    
    try:
        doc = SimpleDocTemplate(temp_file, pagesize=A4,
                                topMargin=inch, bottomMargin=inch,
                                leftMargin=1.2*inch, rightMargin=1.2*inch)
        
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='Justify', alignment=TA_JUSTIFY, fontSize=11, leading=14, fontName='Helvetica'))
        styles.add(ParagraphStyle(name='CenterTitle', alignment=TA_CENTER, fontSize=22, leading=28, spaceAfter=20, fontName='Helvetica-Bold'))
        styles.add(ParagraphStyle(name='Heading1', alignment=TA_LEFT, fontSize=16, leading=20, spaceAfter=12, fontName='Helvetica-Bold'))
        styles.add(ParagraphStyle(name='Heading2', alignment=TA_LEFT, fontSize=14, leading=18, spaceAfter=10, fontName='Helvetica-Bold'))
        
        story = []
        
        # Portada
        story.append(Paragraph(titulo.upper(), styles['CenterTitle']))
        story.append(Spacer(1, 24))
        story.append(Paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}", styles['Normal']))
        story.append(Spacer(1, 12))
        story.append(PageBreak())
        
        # Contenido
        lineas = contenido_markdown.split('\n')
        for linea in lineas:
            if linea.strip():
                if linea.startswith('# '):
                    story.append(Paragraph(linea[2:], styles['Heading1']))
                elif linea.startswith('## '):
                    story.append(Paragraph(linea[3:], styles['Heading2']))
                else:
                    story.append(Paragraph(linea, styles['Justify']))
                story.append(Spacer(1, 6))
        
        doc.build(story)
        return temp_file
    except Exception as e:
        st.error(f"Error generando PDF: {str(e)}")
        return None

# ============================================
# FUNCIONES DE LLAMADAS A API CON REINTENTOS
# ============================================

def llamar_deepseek(prompt, max_intentos=3):
    if not deepseek_client:
        return "⚠️ DeepSeek no configurado"
    for intento in range(max_intentos):
        try:
            response = deepseek_client.chat.completions.create(
                model="deepseek-chat",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2,
                max_tokens=8000
            )
            return response.choices[0].message.content
        except Exception as e:
            if intento == max_intentos - 1:
                return f"❌ Error DeepSeek: {str(e)}"
            time.sleep(2)
    return "❌ Error en llamada a DeepSeek"

def llamar_groq(prompt, max_intentos=3):
    if not groq_client:
        return "⚠️ Groq no configurado"
    for intento in range(max_intentos):
        try:
            response = groq_client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2,
                max_tokens=8000
            )
            return response.choices[0].message.content
        except Exception as e:
            if intento == max_intentos - 1:
                return f"❌ Error Groq: {str(e)}"
            time.sleep(2)
    return "❌ Error en llamada a Groq"

# ============================================
# AGENTE 1: EXTRACCIÓN DE REQUISITOS (CRÍTICO)
# ============================================

def extraer_requisitos_critico(texto):
    """Extrae CADA requisito con especificaciones técnicas detalladas"""
    prompt = f"""
    Eres un AUDITOR POSTDOCTORAL EXPERTO. Extrae CADA requisito de esta convocatoria.
    Actúa como si fueras el ORIGINADOR de la convocatoria.
    
    === CONVOCATORIA ===
    {texto[:12000]}
    
    Responde EXACTAMENTE en este formato JSON. Sé EXTREMADAMENTE DETALLADO y PROFESIONAL:
    
    {{
        "requisitos": [
            {{
                "id": 1,
                "seccion": "sección exacta del documento",
                "texto_original": "texto completo del requisito",
                "especificacion_tecnica": "detalle técnico específico para cumplirlo",
                "plazo": "fecha exacta o 'No especificado'",
                "prioridad": "alta/media/baja",
                "categoria": "administrativo/tecnico/economico/metodologico/etica",
                "es_obligatorio": true,
                "criterio_verificacion": "cómo medir si se cumple este requisito"
            }}
        ],
        "perfiles_requeridos": [
            {{
                "perfil": "nombre del perfil",
                "nivel": "doctorado/maestria/ingenieria/tecnico",
                "experiencia": "años y área específica",
                "habilidades": ["habilidad1", "habilidad2"],
                "dedicacion": "completa/parcial",
                "es_obligatorio": true
            }}
        ],
        "fechas_clave": [
            {{"evento": "apertura/cierre/evaluacion", "fecha": "fecha exacta", "formato": "DD/MM/AAAA"}}
        ],
        "presupuesto_maximo": "monto exacto en USD o 'No especificado'",
        "criterios_evaluacion": [
            {{"criterio": "nombre", "ponderacion": "porcentaje", "descripcion": "detalle"}}
        ],
        "documentos_requeridos": [
            {{"documento": "nombre", "formato": "PDF/Word", "max_paginas": "número o 'N/A'"}}
        ]
    }}
    
    Si un campo no existe en la convocatoria, pon "No especificado" o [].
    """
    return llamar_deepseek(prompt)

# ============================================
# AGENTE 2: VERIFICACIÓN DE CUMPLIMIENTO (100% OBLIGATORIO)
# ============================================

def verificar_cumplimiento_estricto(propuesta, requisitos):
    """Verifica que la propuesta cumpla el 100% de los requisitos. Si no, la rechaza."""
    
    prompt = f"""
    Eres un AUDITOR POSTDOCTORAL ESTRICTÍSIMO. Verifica si la propuesta cumple CADA requisito.
    Actúa como un EVALUADOR OFICIAL de la convocatoria.
    
    === PROPUESTA ===
    {propuesta[:8000]}
    
    === REQUISITOS DE LA CONVOCATORIA ===
    {requisitos[:5000]}
    
    Tu verificación debe ser EXTREMADAMENTE RIGUROSA. No aceptes nada incompleto.
    
    Responde EXACTAMENTE con este formato:
    
    ## RESULTADO DE VERIFICACIÓN
    
    ### REQUISITOS CUMPLIDOS:
    | # | Requisito | Sección donde se cumple | Evidencia en la propuesta |
    |---|-----------|------------------------|---------------------------|
    
    ### REQUISITOS NO CUMPLIDOS:
    | # | Requisito | Qué falta específicamente | Corrección necesaria | Prioridad |
    |---|-----------|--------------------------|---------------------|-----------|
    
    ### CALIFICACIÓN POR SECCIÓN:
    - Portada: X/100
    - Resumen ejecutivo: X/100
    - Introducción y estado del arte: X/100
    - Objetivos: X/100
    - Metodología: X/100
    - Cronograma: X/100
    - Presupuesto: X/100
    - Equipo de trabajo: X/100
    - Resultados esperados: X/100
    - Referencias bibliográficas: X/100
    
    ### CALIFICACIÓN TOTAL: X/100
    
    ### APROBACIÓN: 
    - SI el TOTAL es 100/100 → "APROBADO - LA PROPUESTA ES ACEPTADA"
    - SI el TOTAL es menor a 100/100 → "RECHAZADO - LA PROPUESTA NO CUMPLE LOS REQUISITOS"
    
    ### LISTA DE CORRECCIONES OBLIGATORIAS (solo si es RECHAZADO):
    1. [Corrección específica 1]
    2. [Corrección específica 2]
    ...
    
    ### PRÓXIMOS PASOS:
    - [Instrucciones claras y detalladas para corregir]
    """
    
    return llamar_deepseek(prompt)

# ============================================
# AGENTE 3: CORRECTOR OBLIGATORIO
# ============================================

def corregir_propuesta(propuesta_actual, lista_correcciones, nuevos_documentos=""):
    """Corrige la propuesta basada en las correcciones obligatorias"""
    
    prompt = f"""
    Eres un CORRECTOR POSTDOCTORAL EXPERTO. Debes corregir la propuesta para que cumpla TODOS los requisitos.
    Actúa como un ASESOR SENIOR que garantiza la calidad.
    
    === PROPUESTA ACTUAL ===
    {propuesta_actual[:5000]}
    
    === CORRECCIONES OBLIGATORIAS ===
    {lista_correcciones}
    
    === INFORMACIÓN ADICIONAL ===
    {nuevos_documentos[:2000]}
    
    Genera una VERSIÓN CORREGIDA Y MEJORADA que:
    1. Implementa CADA UNA de las correcciones obligatorias
    2. NO omite ningún requisito
    3. Mantiene el formato profesional académico
    4. Incluye TODAS las secciones requeridas
    5. Mejora la redacción y precisión técnica
    
    Responde SOLO con la propuesta corregida completa, sin comentarios adicionales.
    """
    
    return llamar_groq(prompt)

# ============================================
# AGENTE 4: GENERADOR DE PROPUESTA INICIAL
# ============================================

def generar_propuesta_inicial(requisitos, instrucciones, documentos_previos=""):
    """Genera la primera versión de la propuesta"""
    
    prompt = f"""
    Eres un GENERADOR DE PROPUESTAS POSTDOCTORALES DE ÉLITE.
    Actúa como el MEJOR INVESTIGADOR en tu campo.
    
    === REQUISITOS DE LA CONVOCATORIA ===
    {requisitos[:6000]}
    
    === INSTRUCCIONES ESPECÍFICAS ===
    {instrucciones}
    
    === DOCUMENTOS PREVIOS (CVs, proyectos, artículos) ===
    {documentos_previos[:2000]}
    
    Genera una PROPUESTA POSTDOCTORAL DE ALTA CALIDAD con EXACTAMENTE estas secciones:
    
    1. **PORTADA** (título, autores, institución, fecha, contacto)
    2. **RESUMEN EJECUTIVO** (problema, objetivos, metodología, resultados, presupuesto)
    3. **INTRODUCCIÓN Y ESTADO DEL ARTE** (con citas académicas reales)
    4. **OBJETIVOS** (1 general + mínimo 6 específicos + indicadores)
    5. **METODOLOGÍA DETALLADA** (enfoque, fases, técnicas, instrumentos, análisis)
    6. **CRONOGRAMA DE ACTIVIDADES** (tabla con meses 1 a 36)
    7. **PRESUPUESTO DETALLADO** (tabla con rubros, cantidades, costos)
    8. **EQUIPO DE TRABAJO** (tabla con roles, perfiles, dedicación)
    9. **RESULTADOS ESPERADOS E IMPACTO** (científico, social, económico)
    10. **REFERENCIAS BIBLIOGRÁFICAS** (formato APA, mínimo 15 referencias)
    
    La propuesta debe ser PROFESIONAL, RIGUROSA y LISTA PARA PRESENTAR.
    """
    
    return llamar_groq(prompt)

# ============================================
# AGENTE 5: BÚSQUEDA DE REFERENCIAS (REAL con estructura para API)
# ============================================

def buscar_referencias_academicas(tema):
    """Busca referencias académicas reales (estructura para integrar SerpAPI)"""
    
    # Si tienes API key de SerpAPI, descomenta y usa:
    # SERPAPI_KEY = os.getenv("SERPAPI_KEY", "")
    # if SERPAPI_KEY:
    #     return buscar_referencias_serpapi(tema, SERPAPI_KEY)
    
    prompt = f"""
    Busca REFERENCIAS ACADÉMICAS REALES Y VERIFICABLES sobre: {tema}
    
    Proporciona 15 referencias en formato APA de los ÚLTIMOS 5 AÑOS.
    Incluye para cada una:
    - Autores (apellidos e iniciales)
    - Año de publicación
    - Título completo del artículo/libro
    - Nombre de la revista/editorial
    - Volumen, número, páginas
    - DOI o URL (si está disponible)
    
    Prioriza artículos de:
    - Scopus Q1 y Q2
    - Web of Science
    - Google Scholar (fuentes académicas verificadas)
    """
    return llamar_groq(prompt)

# ============================================
# FUNCIÓN PRINCIPAL CON CICLO DE CALIDAD
# ============================================

def generar_propuesta_con_control_calidad(texto_convocatoria, instrucciones, documentos_previos, max_iteraciones=5):
    """Genera propuesta con ciclos de corrección hasta cumplir 100%"""
    
    historial = []
    
    # Paso 1: Extraer requisitos
    st.write("📋 **Paso 1/5:** Extrayendo requisitos de la convocatoria (DeepSeek)...")
    requisitos_json = extraer_requisitos_critico(texto_convocatoria)
    historial.append({"fase": "extraccion_requisitos", "contenido": requisitos_json[:500]})
    
    # Intentar parsear JSON
    try:
        requisitos_dict = json.loads(requisitos_json)
    except:
        requisitos_dict = {"requisitos": [{"id": 1, "texto_original": requisitos_json[:500]}]}
    
    # Paso 2: Generar propuesta inicial
    st.write("✍️ **Paso 2/5:** Generando propuesta inicial (Groq)...")
    propuesta_actual = generar_propuesta_inicial(requisitos_json, instrucciones, documentos_previos)
    historial.append({"fase": "propuesta_inicial", "contenido": propuesta_actual[:500]})
    
    # Paso 3: Ciclo de verificación y corrección
    st.write("🔄 **Paso 3/5:** Iniciando ciclo de control de calidad...")
    
    for iteracion in range(1, max_iteraciones + 1):
        st.write(f"  - **Iteración {iteracion}:** Verificando cumplimiento...")
        
        # Verificar propuesta actual
        verificacion = verificar_cumplimiento_estricto(propuesta_actual, requisitos_json)
        historial.append({"fase": f"verificacion_iteracion_{iteracion}", "contenido": verificacion[:500]})
        
        # Verificar si está aprobada
        if "APROBADO" in verificacion and "100/100" in verificacion:
            st.success(f"✅ **PROPUESTA APROBADA** después de {iteracion} iteraciones!")
            return propuesta_actual, verificacion, historial, True
        
        # Extraer correcciones obligatorias
        st.write(f"  - Iteración {iteracion}: Propuesta RECHAZADA. Aplicando correcciones...")
        
        # Corregir propuesta
        propuesta_actual = corregir_propuesta(propuesta_actual, verificacion, documentos_previos)
        historial.append({"fase": f"correccion_iteracion_{iteracion}", "contenido": propuesta_actual[:500]})
    
    # Si llegamos aquí, no se aprobó en el máximo de iteraciones
    st.warning(f"⚠️ No se alcanzó el 100% de cumplimiento después de {max_iteraciones} iteraciones.")
    return propuesta_actual, verificacion, historial, False

# ============================================
# INTERFAZ DE STREAMLIT
# ============================================

st.set_page_config(page_title="CrewAI - Postdoctoral Pro", page_icon="🎓", layout="wide")

st.title("🎓 CrewAI - Sistema Postdoctoral de Control de Calidad")
st.markdown("---")

with st.sidebar:
    st.header("📋 Menú")
    opcion = st.radio("Ir a:", ["📝 Nueva Propuesta", "📚 Historial", "⚙️ Estado APIs"])
    
    st.markdown("---")
    st.markdown("### 🎓 Flujo de Control de Calidad")
    st.markdown("""
    1. 📋 **Extraer requisitos** (crítico)
    2. ✍️ **Generar propuesta inicial**
    3. 🔍 **Verificar cumplimiento** (100% obligatorio)
    4. 🔧 **Si NO cumple → Corregir**
    5. 🔁 **Iterar hasta aprobar**
    6. ✅ **Entregar propuesta final**
    """)
    
    st.markdown("---")
    st.markdown("### 📦 Formatos soportados")
    st.markdown("""
    - **Entrada:** PDF, DOCX, TXT, MD
    - **Salida:** Markdown, Word (.docx), PDF (.pdf)
    """)
    
    st.info("⚠️ **Regla estricta:** La propuesta solo se entrega si cumple el 100% de los requisitos.")

if opcion == "⚙️ Estado APIs":
    st.header("🔌 Estado de las APIs y Librerías")
    
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("APIs")
        st.success("✅ DeepSeek") if DEEPSEEK_API_KEY else st.error("❌ DeepSeek")
        st.success("✅ Groq") if GROQ_API_KEY else st.error("❌ Groq")
    
    with col2:
        st.subheader("Librerías")
        st.success("✅ PDF (pypdf)") if PDF_EXTRACTION else st.warning("⚠️ pypdf no instalado")
        st.success("✅ Word (python-docx)") if WORD_AVAILABLE else st.warning("⚠️ python-docx no instalado")
        st.success("✅ PDF (reportlab)") if PDF_AVAILABLE else st.warning("⚠️ reportlab no instalado")

elif opcion == "📝 Nueva Propuesta":
    st.header("🎓 Generación de Propuesta Postdoctoral con Control de Calidad")
    
    with st.form("form_postdoctoral"):
        titulo = st.text_input("📌 Título de la convocatoria / proyecto")
        
        archivos_convocatoria = st.file_uploader(
            "📄 Sube las bases, términos de referencia, requisitos",
            type=["txt", "md", "pdf", "docx"],
            accept_multiple_files=True,
            help="El sistema extraerá TODO el texto de estos documentos"
        )
        
        archivos_previos = st.file_uploader(
            "📎 Sube documentos adicionales (CVs, artículos, proyectos previos)",
            type=["txt", "md", "pdf", "docx"],
            accept_multiple_files=True,
            help="Opcional: para enriquecer la propuesta"
        )
        
        instrucciones = st.text_area(
            "📝 Instrucciones específicas",
            height=120,
            placeholder="""Ejemplo:
            - Duración del proyecto: 36 meses
            - Presupuesto máximo: $250,000 USD
            - Equipo mínimo: 1 Investigador Principal + 2 Postdocs + 1 Estudiante doctoral
            - Incluir análisis estadístico con R o Python
            - Priorizar metodologías mixtas
            - Publicaciones esperadas: mínimo 3 artículos Q1
            """,
            help="Cuantos más detalles, mejor será la propuesta"
        )
        
        formato_salida = st.selectbox(
            "📁 Formato de salida",
            ["Markdown (.md)", "Word (.docx)", "PDF (.pdf)"],
            help="Selecciona el formato en el que quieres recibir la propuesta final"
        )
        
        max_iteraciones = st.slider(
            "🔄 Máximo de iteraciones de corrección",
            min_value=1,
            max_value=10,
            value=5,
            help="Número máximo de ciclos de corrección antes de entregar la última versión"
        )
        
        submitted = st.form_submit_button("🚀 INICIAR PROCESO DE CALIDAD", type="primary")
        
        if submitted:
            if not titulo:
                st.error("❌ Ingresa un título")
            elif not archivos_convocatoria:
                st.error("❌ Sube al menos un documento de la convocatoria")
            else:
                # Extraer texto de archivos
                with st.spinner("📄 Extrayendo texto de los documentos..."):
                    texto_convocatoria = extraer_texto_multiple(archivos_convocatoria)
                    texto_previos = extraer_texto_multiple(archivos_previos) if archivos_previos else ""
                
                # Ejecutar proceso con control de calidad
                propuesta_final, verificacion_final, historial, aprobada = generar_propuesta_con_control_calidad(
                    texto_convocatoria,
                    instrucciones,
                    texto_previos,
                    max_iteraciones
                )
                
                # Guardar en base de datos
                guardar_analisis(
                    titulo,
                    instrucciones,
                    formato_salida,
                    "APA 7ª edición",
                    propuesta_final,
                    [a.name for a in archivos_convocatoria]
                )
                
                # Mostrar resultados
                st.markdown("---")
                
                if aprobada:
                    st.balloons()
                    st.success("🎉 **¡PROPUESTA APROBADA!**")
                    st.success("Cumple con el 100% de los requisitos de la convocatoria.")
                else:
                    st.warning("⚠️ **PROPUESTA CON CORRECCIONES PENDIENTES**")
                    st.info("Se alcanzó el máximo de iteraciones. Revise las correcciones sugeridas.")
                
                st.markdown("---")
                st.header("📄 PROPUESTA FINAL")
                st.markdown(propuesta_final)
                
                # Botones de descarga según formato
                if formato_salida == "Word (.docx)":
                    doc_file = generar_documento_word(propuesta_final, titulo)
                    if doc_file:
                        with open(doc_file, "rb") as f:
                            st.download_button("📥 DESCARGAR WORD", f, file_name=doc_file, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                elif formato_salida == "PDF (.pdf)":
                    pdf_file = generar_documento_pdf(propuesta_final, titulo)
                    if pdf_file:
                        with open(pdf_file, "rb") as f:
                            st.download_button("📥 DESCARGAR PDF", f, file_name=pdf_file, mime="application/pdf")
                else:
                    st.download_button("📥 DESCARGAR MARKDOWN", propuesta_final, file_name=f"{titulo.replace(' ', '_')}.md")
                
                with st.expander("📋 VERIFICACIÓN FINAL"):
                    st.markdown(verificacion_final)
                
                with st.expander("🔄 HISTORIAL DE ITERACIONES"):
                    for i, item in enumerate(historial):
                        st.write(f"**{i+1}. {item['fase']}**")
                        st.text(item['contenido'][:500] + "...")
                        st.markdown("---")

else:
    st.header("📚 Historial de Propuestas")
    
    lista = obtener_analisis()
    
    if not lista:
        st.info("No hay propuestas previas. Genera una nueva propuesta postdoctoral.")
    else:
        for item in lista:
            with st.expander(f"🎓 {item.get('titulo', 'Sin título')} - {item.get('fecha', '')[:10]}"):
                st.write(f"**Formato:** {item.get('formato_salida', 'N/A')}")
                st.write(f"**Estilo citas:** {item.get('estilo_citas', 'N/A')}")
                with st.expander("Ver propuesta"):
                    st.markdown(item.get('resultado', 'No disponible')[:3000])

st.markdown("---")
st.caption("🎓 CrewAI Postdoctoral Pro | DeepSeek + Groq | Control de calidad 100% | Iteración hasta aprobación | Word/PDF real")